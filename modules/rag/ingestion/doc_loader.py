"""
Ingestion pipeline for 3GPP spec .docx / .doc files.
Uses python-docx for .docx and mammoth for legacy OLE2 .doc files.
Chunks text and ingests into Weaviate.
"""

import os
import re
from pathlib import Path
from typing import List

import docx
import mammoth
from loguru import logger
from tqdm import tqdm

from modules.rag.vector_store.schema import get_client, COLLECTION_NAME
from modules.rag.vector_store.retriever import embed

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "512"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "64"))

SPEC_ID_MAP = {
    "28532": "TS_28.532",
    "28552": "TS_28.552",
    "32111": "TS_32.111",
}


def _detect_spec_id(filename: str) -> str:
    for key, spec_id in SPEC_ID_MAP.items():
        if key in filename:
            return spec_id
    return "UNKNOWN"


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunks.append(" ".join(words[start:end]))
        start += chunk_size - overlap
    return [c for c in chunks if len(c.strip()) > 20]


def _extract_section_heading(text: str) -> str:
    match = re.match(r"^(\d+(\.\d+)*\s+[A-Z][^\n]{5,60})", text.strip())
    return match.group(1).strip() if match else ""


def _read_docx(path: Path) -> str:
    """Extract full text from .docx file."""
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_doc(path: Path) -> str:
    """Extract text from legacy OLE2 .doc file using mammoth."""
    try:
        with open(path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        return result.value
    except Exception as e:
        logger.warning(f"mammoth failed for {path.name}: {e} — skipping")
        return ""


def ingest_doc(doc_path: str | Path) -> int:
    """
    Ingest a single .docx or .doc file into Weaviate.
    Returns number of chunks inserted.
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"File not found: {doc_path}")

    spec_id = _detect_spec_id(doc_path.name)
    source = doc_path.name
    suffix = doc_path.suffix.lower()

    logger.info(f"Ingesting {source} (spec: {spec_id})")

    if suffix == ".docx":
        raw_text = _read_docx(doc_path)
    elif suffix == ".doc":
        raw_text = _read_doc(doc_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    if not raw_text.strip():
        logger.warning(f"No text extracted from {source}")
        return 0

    chunks = _chunk_text(raw_text)
    objects_to_insert = []
    for i, chunk in enumerate(tqdm(chunks, desc=f"Embedding {source}", unit="chunk")):
        vector = embed(chunk)
        objects_to_insert.append({
            "text": chunk,
            "source": source,
            "doc_type": "docx",
            "spec_id": spec_id,
            "page_number": 0,
            "chunk_index": i,
            "section": _extract_section_heading(chunk),
            "_vector": vector,
        })

    inserted = 0
    with get_client() as client:
        collection = client.collections.get(COLLECTION_NAME)
        with collection.batch.dynamic() as batch:
            for obj in objects_to_insert:
                vector = obj.pop("_vector")
                batch.add_object(properties=obj, vector=vector)
                inserted += 1

    logger.success(f"Inserted {inserted} chunks from {source}")
    return inserted


def ingest_vocabulary_docx(docx_path: str | Path) -> int:
    """
    Ingest the 3GPP_vocabulary.docx as additional knowledge base context.
    Stored with doc_type='vocabulary' so it can be filtered separately.
    """
    doc_path = Path(docx_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Vocabulary file not found: {doc_path}")

    logger.info(f"Ingesting vocabulary: {doc_path.name}")
    doc = docx.Document(str(doc_path))
    raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chunks = _chunk_text(raw_text)

    inserted = 0
    with get_client() as client:
        collection = client.collections.get(COLLECTION_NAME)
        with collection.batch.dynamic() as batch:
            for i, chunk in enumerate(tqdm(chunks, desc=f"Embedding vocabulary", unit="chunk")):
                vector = embed(chunk)
                batch.add_object(
                    properties={
                        "text": chunk,
                        "source": doc_path.name,
                        "doc_type": "vocabulary",
                        "spec_id": "3GPP_vocabulary",
                        "page_number": 0,
                        "chunk_index": i,
                        "section": _extract_section_heading(chunk),
                    },
                    vector=vector,
                )
                inserted += 1

    logger.success(f"Inserted {inserted} vocabulary chunks from {doc_path.name}")
    return inserted


def ingest_all_docs(doc_dir: str | Path) -> None:
    """Ingest all .docx and .doc files in the given directory."""
    doc_dir = Path(doc_dir)
    docs = list(doc_dir.glob("*.docx")) + list(doc_dir.glob("*.doc"))
    if not docs:
        logger.warning(f"No .docx/.doc files found in {doc_dir}")
        return
    for doc in docs:
        try:
            ingest_doc(doc)
        except Exception as e:
            logger.error(f"Failed to ingest {doc.name}: {e}")


if __name__ == "__main__":
    ingest_all_docs("data/raw/3gpp_specs")
    ingest_vocabulary_docx("data/raw/telecom_complaints/3GPP_vocabulary.docx")
